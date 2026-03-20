#!/usr/bin/env python3
"""CV renderer: generates PDF and Word documents from cv_data.yaml."""

import argparse
import sys
from pathlib import Path

import yaml
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib.colors import HexColor
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    HRFlowable,
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def load_cv(path: str = "cv_data.yaml") -> dict:
    with open(path, "r", encoding="utf-8") as f:
        return yaml.safe_load(f)


# ---------------------------------------------------------------------------
# PDF rendering (ReportLab)
# ---------------------------------------------------------------------------

def _pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        "CVName", parent=styles["Title"], fontSize=20, leading=24,
        spaceAfter=2, textColor=HexColor("#1a1a1a"),
    ))
    styles.add(ParagraphStyle(
        "CVContact", parent=styles["Normal"], fontSize=10, leading=14,
        alignment=1, spaceAfter=10, textColor=HexColor("#444444"),
    ))
    styles.add(ParagraphStyle(
        "CVSection", parent=styles["Heading2"], fontSize=13, leading=16,
        spaceBefore=14, spaceAfter=4, textColor=HexColor("#0b3d91"),
    ))
    styles.add(ParagraphStyle(
        "CVBody", parent=styles["Normal"], fontSize=10, leading=14,
        spaceAfter=2,
    ))
    styles.add(ParagraphStyle(
        "CVBodyBold", parent=styles["Normal"], fontSize=10, leading=14,
        spaceAfter=2,
    ))
    styles.add(ParagraphStyle(
        "CVBodyItalic", parent=styles["Normal"], fontSize=10, leading=14,
        spaceAfter=2, textColor=HexColor("#555555"),
    ))
    return styles


def _add_section_hr(story):
    story.append(HRFlowable(
        width="100%", thickness=0.5, color=HexColor("#cccccc"),
        spaceAfter=6, spaceBefore=0,
    ))


def render_pdf(cv: dict, output: str = "cv_output.pdf"):
    doc = SimpleDocTemplate(
        output, pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm,
        topMargin=15 * mm, bottomMargin=15 * mm,
    )
    s = _pdf_styles()
    story: list = []

    # Header
    story.append(Paragraph(cv["name"], s["CVName"]))
    contact = f'{cv["email"]} | {cv["phone"]}'
    story.append(Paragraph(contact, s["CVContact"]))

    # Research Interests
    story.append(Paragraph("RESEARCH INTERESTS", s["CVSection"]))
    _add_section_hr(story)
    story.append(Paragraph(", ".join(cv["research_interests"]), s["CVBody"]))

    # Education
    story.append(Paragraph("EDUCATION", s["CVSection"]))
    _add_section_hr(story)
    for e in cv["education"]:
        story.append(Paragraph(
            f'<b>{e["degree"]}</b> — {e["institution"]}', s["CVBody"],
        ))
        parts = [e["period"]]
        if e.get("advisor"):
            parts.append(f'Advisor: {e["advisor"]}')
        parts.append(e["location"])
        story.append(Paragraph(" | ".join(parts), s["CVBodyItalic"]))
        story.append(Spacer(1, 4))

    # Publications
    if cv.get("publications"):
        story.append(Paragraph("PUBLICATIONS", s["CVSection"]))
        _add_section_hr(story)
        for i, p in enumerate(cv["publications"], 1):
            story.append(Paragraph(
                f'[{i}] <b>{p["title"]}</b>', s["CVBody"],
            ))
            story.append(Paragraph(p["authors"], s["CVBodyItalic"]))
            venue = p["venue"]
            if p.get("note"):
                venue += f'  {p["note"]}'
            story.append(Paragraph(venue, s["CVBodyItalic"]))
            story.append(Spacer(1, 4))

    # Conferences
    if cv.get("conferences"):
        story.append(Paragraph("CONFERENCES", s["CVSection"]))
        _add_section_hr(story)
        for i, c in enumerate(cv["conferences"], 1):
            story.append(Paragraph(
                f'[{i}] <b>{c["title"]}</b>', s["CVBody"],
            ))
            story.append(Paragraph(c["venue"], s["CVBodyItalic"]))
            story.append(Paragraph(c["authors"], s["CVBody"]))
            story.append(Spacer(1, 4))

    # Work Experience
    if cv.get("work_experience"):
        story.append(Paragraph("WORK EXPERIENCE", s["CVSection"]))
        _add_section_hr(story)
        for w in cv["work_experience"]:
            story.append(Paragraph(
                f'<b>{w["position"]}</b> — {w["organization"]}', s["CVBody"],
            ))
            story.append(Paragraph(
                f'{w["period"]} | {w["location"]}', s["CVBodyItalic"],
            ))
            story.append(Spacer(1, 4))

    # Research Experience
    if cv.get("research_experience"):
        story.append(Paragraph("RESEARCH EXPERIENCE", s["CVSection"]))
        _add_section_hr(story)
        for r in cv["research_experience"]:
            title = f'<b>{r["position"]}</b> — {r["organization"]}'
            if r.get("advisor"):
                title += f' (Advisor: {r["advisor"]})'
            story.append(Paragraph(title, s["CVBody"]))
            story.append(Paragraph(
                f'{r["period"]} | {r["location"]}', s["CVBodyItalic"],
            ))
            if r.get("subject"):
                story.append(Paragraph(
                    f'Subject: {r["subject"]}', s["CVBody"],
                ))
            story.append(Spacer(1, 4))

    doc.build(story)
    print(f"PDF saved: {output}")


# ---------------------------------------------------------------------------
# Word rendering (python-docx)
# ---------------------------------------------------------------------------

def _add_heading_with_line(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x0B, 0x3D, 0x91)
    p.space_after = Pt(2)
    # thin line
    border_p = doc.add_paragraph()
    border_p.space_before = Pt(0)
    border_p.space_after = Pt(6)
    pPr = border_p._element.get_or_add_pPr()
    from docx.oxml.ns import qn
    from lxml import etree
    pBdr = etree.SubElement(pPr, qn("w:pBdr"))
    bottom = etree.SubElement(pBdr, qn("w:bottom"))
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")


def _add_entry(doc: Document, bold_text: str, italic_text: str = "", body_text: str = ""):
    if bold_text:
        p = doc.add_paragraph()
        run = p.add_run(bold_text)
        run.bold = True
        run.font.size = Pt(10)
        p.space_after = Pt(1)
    if italic_text:
        p = doc.add_paragraph()
        run = p.add_run(italic_text)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        p.space_after = Pt(1)
    if body_text:
        p = doc.add_paragraph()
        run = p.add_run(body_text)
        run.font.size = Pt(10)
        p.space_after = Pt(1)


def render_word(cv: dict, output: str = "cv_output.docx"):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(cv["name"])
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    p.space_after = Pt(2)

    # Contact
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{cv["email"]} | {cv["phone"]}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    p.space_after = Pt(10)

    # Research Interests
    _add_heading_with_line(doc, "RESEARCH INTERESTS")
    p = doc.add_paragraph()
    run = p.add_run(", ".join(cv["research_interests"]))
    run.font.size = Pt(10)

    # Education
    _add_heading_with_line(doc, "EDUCATION")
    for e in cv["education"]:
        parts = [e["period"]]
        if e.get("advisor"):
            parts.append(f'Advisor: {e["advisor"]}')
        parts.append(e["location"])
        _add_entry(
            doc,
            f'{e["degree"]} — {e["institution"]}',
            " | ".join(parts),
        )

    # Publications
    if cv.get("publications"):
        _add_heading_with_line(doc, "PUBLICATIONS")
        for i, pub in enumerate(cv["publications"], 1):
            venue = pub["venue"]
            if pub.get("note"):
                venue += f'  {pub["note"]}'
            _add_entry(doc, f'[{i}] {pub["title"]}', pub["authors"], venue)

    # Conferences
    if cv.get("conferences"):
        _add_heading_with_line(doc, "CONFERENCES")
        for i, c in enumerate(cv["conferences"], 1):
            _add_entry(doc, f'[{i}] {c["title"]}', c["venue"], c["authors"])

    # Work Experience
    if cv.get("work_experience"):
        _add_heading_with_line(doc, "WORK EXPERIENCE")
        for w in cv["work_experience"]:
            _add_entry(
                doc,
                f'{w["position"]} — {w["organization"]}',
                f'{w["period"]} | {w["location"]}',
            )

    # Research Experience
    if cv.get("research_experience"):
        _add_heading_with_line(doc, "RESEARCH EXPERIENCE")
        for r in cv["research_experience"]:
            title = f'{r["position"]} — {r["organization"]}'
            if r.get("advisor"):
                title += f' (Advisor: {r["advisor"]})'
            sub = f'{r["period"]} | {r["location"]}'
            body = f'Subject: {r["subject"]}' if r.get("subject") else ""
            _add_entry(doc, title, sub, body)

    doc.save(output)
    print(f"Word saved: {output}")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Render CV to PDF and/or Word")
    parser.add_argument("--data", default="cv_data.yaml", help="Path to CV YAML data")
    parser.add_argument("--format", choices=["pdf", "word", "both"], default="both",
                        help="Output format (default: both)")
    parser.add_argument("--pdf-output", default="cv_output.pdf", help="PDF output filename")
    parser.add_argument("--word-output", default="cv_output.docx", help="Word output filename")
    args = parser.parse_args()

    cv = load_cv(args.data)

    if args.format in ("pdf", "both"):
        render_pdf(cv, args.pdf_output)
    if args.format in ("word", "both"):
        render_word(cv, args.word_output)


if __name__ == "__main__":
    main()
